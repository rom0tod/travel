"""Сборка Word-версий ТЗ и пояснительной записки.

Конвертирует исходники из ``docs/*.md`` в .docx с аккуратной типографией:
заголовки, абзацы, списки, моноширинные блоки, рамки таблиц.

Запуск:
    python -m tools.build_docs
"""
from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


HEADER_FONT = "Calibri"
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

ACCENT = RGBColor(0x25, 0x63, 0xEB)
NAVY = RGBColor(0x0F, 0x17, 0x2A)
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x64, 0x74, 0x8B)


# --- Стили --------------------------------------------------------------

def configure_document(doc: Document) -> None:
    """Поля A4, базовые шрифты в стилях по умолчанию."""
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = HEADER_FONT
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(8)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(20)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = HEADER_FONT
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = HEADER_FONT
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = HEADER_FONT
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def add_paragraph(doc: Document, text: str) -> None:
    """Параграф с поддержкой *италика* и **жирного** через markdown."""
    p = doc.add_paragraph()
    _add_inline_runs(p, text)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    _add_inline_runs(p, text)
    p.paragraph_format.space_after = Pt(2)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    _add_inline_runs(p, text)
    p.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, code: str) -> None:
    """Моноширинный блок с серой рамкой-фоном."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    run.font.name = MONO_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY

    # Левая полоса-акцент через нижний бордер абзаца
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2563EB")
    pBdr.append(left)
    pPr.append(pBdr)


def add_table(doc: Document, header: List[str],
              rows: List[List[str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = "Light Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for index, label in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.size = Pt(11)
        # Голубая заливка
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "DBEAFE")
        tcPr.append(shd)

    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            cell = table.rows[row_index].cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(10.5)
            run.font.color.rgb = TEXT


def _add_inline_runs(paragraph, text: str) -> None:
    r"""Парсер `**bold**`, `*italic*` и `\`code\`` для одной строки."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run.text = part[1:-1]
            run.font.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = MONO_FONT
            run.font.size = Pt(10)
        else:
            run.text = part


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CBD5E1")
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# --- Документы ---------------------------------------------------------------

@dataclass
class Author:
    fio: str = "Роман Тодор"
    course: str = "Яндекс Лицей · «Основы промышленного программирования»"
    module: str = "WebServer + REST API"
    year: str = "2026"


def build_tz(target: Path, author: Author) -> None:
    doc = Document()
    configure_document(doc)

    add_title(doc, "Техническое задание")
    add_subtitle(
        doc,
        "Проект «TripPlanner» — социальный планировщик путешествий",
    )

    add_paragraph(doc, f"**Автор:** {author.fio}")
    add_paragraph(doc, f"**Курс:** {author.course}")
    add_paragraph(doc, f"**Модуль:** {author.module}")
    add_paragraph(doc, "**Тип проекта:** веб-приложение (Flask)")
    add_horizontal_rule(doc)

    add_h1(doc, "1. Назначение")
    add_paragraph(
        doc,
        "TripPlanner — социальный планировщик путешествий, позволяющий "
        "пользователю собрать маршрут поездки по дням с привязкой к "
        "карте, вести бюджет и делиться готовыми поездками с "
        "сообществом.",
    )

    add_h1(doc, "2. Целевая аудитория")
    add_bullet(doc, "Студенты и школьники, планирующие самостоятельные поездки.")
    add_bullet(
        doc,
        "Путешественники, которые хотят сохранить и переиспользовать "
        "свои маршруты, а также вдохновляться чужими.",
    )

    add_h1(doc, "3. Функциональные требования")

    add_h2(doc, "3.1. Учётные записи")
    add_bullet(doc, "Регистрация по логину и e-mail с проверкой уникальности.")
    add_bullet(doc, "Хэширование паролей (`werkzeug.security`).")
    add_bullet(doc, "Вход с опцией «запомнить меня».")
    add_bullet(doc, "Публичный профиль пользователя с аватаром и описанием.")

    add_h2(doc, "3.2. Поездки")
    add_bullet(doc, "Создание / редактирование / удаление поездки.")
    add_bullet(
        doc,
        "Поля: название, направление, описание, даты, бюджет, обложка, "
        "теги, флаг публичности.",
    )
    add_bullet(
        doc,
        "Просмотр поездки доступен владельцу всегда, остальным — только "
        "если поездка публичная.",
    )

    add_h2(doc, "3.3. Маршрут")
    add_bullet(doc, "Добавление точек интереса с координатами (широта/долгота).")
    add_bullet(
        doc,
        "Автогеокодирование адреса через сторонний API "
        "Nominatim (OpenStreetMap).",
    )
    add_bullet(doc, "Распределение точек по дням поездки.")
    add_bullet(
        doc,
        "Шесть категорий точек: достопримечательность, жильё, еда, "
        "транспорт, активность, прочее.",
    )
    add_bullet(
        doc,
        "Отображение всех точек на интерактивной карте Leaflet.",
    )

    add_h2(doc, "3.4. Бюджет")
    add_bullet(doc, "Учёт расходов по 6 категориям и в произвольной валюте.")
    add_bullet(doc, "Автоматический подсчёт суммы трат и остатка от бюджета.")
    add_bullet(doc, "Индикатор перерасхода.")

    add_h2(doc, "3.5. Социальные функции")
    add_bullet(doc, "Лайк публичной поездки (AJAX без перезагрузки).")
    add_bullet(doc, "Комментарии: оставить, удалить (автор или владелец).")
    add_bullet(
        doc,
        "Explore-страница: поиск по тексту, фильтр по тегам, пагинация.",
    )

    add_h2(doc, "3.6. REST API")
    add_bullet(
        doc,
        "Полный CRUD поездок: `GET/POST /api/trips`, "
        "`GET/PUT/DELETE /api/trips/<id>`.",
    )
    add_bullet(doc, "CRUD точек маршрута: `/api/trips/<id>/places[/<pid>]`.")
    add_bullet(doc, "Расходы: `/api/trips/<id>/expenses[/<eid>]`.")
    add_bullet(
        doc,
        "Пользователи (read-only): `/api/users`, `/api/users/<id>`.",
    )
    add_bullet(doc, "Прокси-эндпоинт геокодера: `/api/geocode?address=...`.")
    add_bullet(
        doc,
        "Формат ответа: `{\"success\": true, ...}` или "
        "`{\"success\": false, \"error\": \"...\"}`.",
    )

    add_h1(doc, "4. Нефункциональные требования")
    add_table(
        doc,
        ["Свойство", "Требование"],
        [
            ["Время отклика", "< 300 мс на типовой запрос на локальной БД"],
            ["Безопасность", "Хэшированные пароли, CSRF-токены на всех POST-формах"],
            ["Лимит файла", "До 8 МБ на загрузку (MAX_CONTENT_LENGTH)"],
            ["Совместимость браузеров", "Современные Chrome / Firefox / Edge"],
            ["Мобильный режим", "Адаптивная вёрстка через Bootstrap 5"],
        ],
    )

    add_h1(doc, "5. Используемый стек")
    add_table(
        doc,
        ["Слой", "Технология"],
        [
            ["Бэкенд", "Python 3.12, Flask 2.3+"],
            ["ORM", "SQLAlchemy 2.0 + sqlalchemy-serializer"],
            ["Аутентификация", "Flask-Login"],
            ["Формы", "Flask-WTF + WTForms + email-validator"],
            ["REST API", "Flask-RESTful"],
            ["Шаблоны", "Jinja2"],
            ["Фронтенд", "Bootstrap 5, Bootstrap Icons, Leaflet 1.9"],
            ["Сторонний API", "Nominatim (OpenStreetMap)"],
            ["База данных", "SQLite (dev) / PostgreSQL (prod)"],
            ["WSGI", "gunicorn (prod)"],
            ["Хостинг", "Render.com или PythonAnywhere"],
        ],
    )

    add_h1(doc, "6. Структура хранилища")
    add_paragraph(doc, "7 ORM-моделей:")
    add_bullet(doc, "**User** — учётные записи.")
    add_bullet(doc, "**Trip** — поездка (FK → User).")
    add_bullet(doc, "**Place** — точка маршрута (FK → Trip).")
    add_bullet(doc, "**Expense** — расход (FK → Trip).")
    add_bullet(doc, "**Comment** — комментарий (FK → Trip, FK → User).")
    add_bullet(doc, "**Like** — связь «лайк» (FK → Trip, FK → User; UNIQUE).")
    add_bullet(doc, "**Tag** — тег (many-to-many → Trip через trip_tags).")

    add_h1(doc, "7. Внешние интерфейсы")
    add_bullet(doc, "**Веб-страницы** — рендеринг Jinja2.")
    add_bullet(doc, "**REST API** — формат JSON.")
    add_bullet(
        doc,
        "**Геокодер Nominatim** — `GET https://nominatim.openstreetmap.org"
        "/search` с обязательным User-Agent.",
    )

    add_h1(doc, "8. Этапы реализации")
    add_table(
        doc,
        ["№", "Этап", "Артефакт"],
        [
            ["1", "Описание проекта, согласование с преподавателем",
             "README.md"],
            ["2", "ТЗ", "docs/ТЗ.docx (этот документ)"],
            ["3", "Скелет приложения + аутентификация",
             "main.py, data/users.py, forms/auth.py"],
            ["4", "CRUD поездок и точек, карта Leaflet",
             "data/trips.py, places.py, templates/trips/"],
            ["5", "Бюджет, теги, социалка",
             "data/expenses.py, likes.py, comments.py, tags.py"],
            ["6", "REST API + геокодер + черновики документов",
             "api/, tools/geocoder.py"],
            ["7", "Деплой, пояснительная записка, презентация",
             "DEPLOY.md, docs/*"],
        ],
    )

    add_h1(doc, "9. Критерии приёмки")
    add_bullet(
        doc,
        "Все маршруты возвращают HTTP 200/302 для авторизованного "
        "пользователя.",
    )
    add_bullet(
        doc,
        "API возвращает корректный JSON по всем GET-эндпоинтам без "
        "авторизации.",
    )
    add_bullet(
        doc,
        "При попытке доступа к чужой приватной поездке возвращается 403.",
    )
    add_bullet(doc, "Геокодер возвращает координаты для адреса «Москва».")
    add_bullet(
        doc,
        "Демо-сидер создаёт минимум 3 поездки, 6 точек, 3 пользователей.",
    )

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    print(f"ТЗ сохранено: {target}")


def build_explanatory_note(target: Path, author: Author) -> None:
    doc = Document()
    configure_document(doc)

    add_title(doc, "Пояснительная записка")
    add_subtitle(
        doc,
        "Проект «TripPlanner» — социальный планировщик путешествий",
    )

    add_paragraph(doc, f"**Автор:** {author.fio}")
    add_paragraph(doc, f"**Курс:** {author.course}")
    add_paragraph(doc, f"**Модуль:** {author.module}")
    add_paragraph(
        doc,
        "**Репозиторий:** https://github.com/rom0tod/travel",
    )
    add_paragraph(
        doc,
        "**Развёрнутое приложение:** https://tripplanner-17ou.onrender.com",
    )
    add_horizontal_rule(doc)

    add_h1(doc, "1. Введение")

    add_h2(doc, "Идея проекта")
    add_paragraph(
        doc,
        "Большинство планировщиков путешествий, доступных школьнику и "
        "студенту, либо платные, либо привязаны к экосистеме (Google "
        "Maps, Maps.me, Booking). Хочется отдельный «личный» инструмент:",
    )
    add_bullet(doc, "собрать поездку из точек, разложенных по дням;")
    add_bullet(doc, "видеть маршрут на карте;")
    add_bullet(doc, "следить за бюджетом;")
    add_bullet(
        doc,
        "легко делиться готовым маршрутом с друзьями.",
    )
    add_paragraph(doc, "TripPlanner закрывает именно эту нишу.")

    add_h2(doc, "Какие задачи решает")
    add_numbered(
        doc,
        "Снимает с пользователя необходимость держать маршрут «в голове» "
        "или в плоской заметке.",
    )
    add_numbered(
        doc,
        "Даёт визуальное представление поездки (карта Leaflet, "
        "accordion по дням).",
    )
    add_numbered(
        doc,
        "Помогает не выйти за бюджет и видеть распределение расходов "
        "по категориям.",
    )
    add_numbered(
        doc,
        "Превращает приватный маршрут в публикацию для других "
        "путешественников: лайки, комментарии, поиск.",
    )

    add_h1(doc, "2. Архитектура и проектирование")

    add_h2(doc, "2.1. Принципы")
    add_bullet(
        doc,
        "**Слойная архитектура.** Код разделён на изолированные "
        "ответственности: `data/` — модели и сессия БД, `forms/` — "
        "валидация ввода, `api/` — REST-ресурсы, `tools/` — утилиты, "
        "`templates/` + `static/` — представление.",
    )
    add_bullet(
        doc,
        "**Фабрика приложения.** `create_app()` в `main.py` позволяет "
        "создавать конфигурации под dev/prod и упрощает тестирование.",
    )
    add_bullet(
        doc,
        "**Слой данных независим от Flask.** Модели можно использовать "
        "в скриптах сидера без поднятия HTTP-сервера.",
    )
    add_bullet(
        doc,
        "**Конфигурация в одном месте.** Все лимиты, пути и константы — "
        "в `config.py` и завязаны на переменные окружения.",
    )

    add_h2(doc, "2.2. ORM-модель данных")
    add_paragraph(doc, "Семь сущностей с естественными связями:")
    add_code_block(
        doc,
        "User ───< Trip ───< Place\n"
        "       │       └─< Expense\n"
        "       │       └─< Comment >─── User\n"
        "       │       └─< Like    >─── User\n"
        "       │       └─>< Tag   (many-to-many)\n"
        "       └────────< Comment\n"
        "       └────────< Like",
    )
    add_paragraph(
        doc,
        "Все модели наследуются от `SqlAlchemyBase` и подмешивают "
        "`SerializerMixin` для единого `to_dict()` в REST API.",
    )

    add_h2(doc, "2.3. Сценарий «создать поездку с точкой на карте»")
    add_numbered(
        doc,
        "Пользователь заполняет форму поездки (`forms/trip.py`, WTForms "
        "+ CSRF от Flask-WTF).",
    )
    add_numbered(
        doc,
        "View-функция `create_trip` сохраняет объект в БД через сессию "
        "SQLAlchemy.",
    )
    add_numbered(
        doc,
        "На странице деталей пользователь нажимает «Добавить точку».",
    )
    add_numbered(
        doc,
        "На странице точки JS-функция `initPlacePicker` поднимает карту "
        "Leaflet. Клик по карте обновляет hidden-инпуты с координатами.",
    )
    add_numbered(
        doc,
        "Альтернативно пользователь вводит адрес, JS делает "
        "`GET /api/geocode?address=…`. Серверный эндпоинт обращается "
        "к Nominatim, возвращает JSON с координатами.",
    )
    add_numbered(
        doc,
        "Сохранение точки приводит к обновлению карты и расписания.",
    )

    add_h2(doc, "2.4. Безопасность")
    add_table(
        doc,
        ["Угроза", "Защита"],
        [
            ["Подсмотренный пароль в БД",
             "Хэширование werkzeug.security"],
            ["CSRF на POST-формах",
             "Токен flask_wtf.FlaskForm"],
            ["Просмотр чужой приватной поездки",
             "Метод Trip.is_visible_to(user)"],
            ["Изменение чужой поездки",
             "Trip.is_editable_by(user) во view и API"],
            ["Загрузка опасного файла",
             "Проверка расширения + secure_filename, лимит 8 МБ"],
            ["Перебор пароля",
             "Сообщение об ошибке не уточняет, что неверно"],
        ],
    )

    add_h2(doc, "2.5. Внешние интеграции")
    add_bullet(
        doc,
        "**OpenStreetMap (тайлы)** — без ключа, через CDN.",
    )
    add_bullet(
        doc,
        "**Nominatim** — поиск координат по адресу, обязательно с "
        "собственным User-Agent (требование сервиса).",
    )

    add_h1(doc, "3. Применённые технологии")
    add_table(
        doc,
        ["Категория", "Технологии и реализация"],
        [
            ["Управление зависимостями", "requirements.txt"],
            ["Веб-фреймворк", "Flask с фабрикой приложения"],
            ["ORM", "SQLAlchemy 2.0, DeclarativeBase"],
            ["Сериализация моделей", "sqlalchemy_serializer.SerializerMixin"],
            ["Аутентификация", "Flask-Login, хэш паролей, remember-me"],
            ["Формы", "Flask-WTF + WTForms, кастомные валидаторы"],
            ["Файлы", "Загрузка обложек и аватаров"],
            ["REST API", "Flask-RESTful, 8 ресурсов CRUD"],
            ["Сторонний API", "Nominatim (геокодирование адресов)"],
            ["Шаблоны", "Jinja2 с наследованием, фильтры, контекст"],
            ["Frontend", "Bootstrap 5, Leaflet 1.9, AJAX через fetch"],
            ["Хранение", "SQLite (dev), PostgreSQL (prod)"],
            ["Хостинг", "Render.com (Blueprint render.yaml)"],
            ["Production WSGI", "gunicorn"],
        ],
    )

    add_h1(doc, "4. Что нового по сравнению с типовыми проектами")
    add_bullet(
        doc,
        "**Интерактивная карта Leaflet** с маршрутами по дням и выбором "
        "координат кликом.",
    )
    add_bullet(
        doc,
        "**Геокодер на стороннем API** с обработкой ошибок сети.",
    )
    add_bullet(
        doc,
        "**Многоуровневое разделение прав:** публичная/приватная "
        "поездка, владелец может удалить чужой комментарий под своей "
        "публикацией.",
    )
    add_bullet(
        doc,
        "**Постатейный учёт бюджета** с сводкой по категориям.",
    )
    add_bullet(
        doc,
        "**Чистая REST-обёртка**: одинаковый формат ответа, явные коды "
        "состояния, поддержка пагинации.",
    )
    add_bullet(
        doc,
        "**Production-ready деплой**: Procfile + render.yaml + wsgi.py "
        "+ чтение конфигурации из переменных окружения.",
    )

    add_h1(doc, "5. Тестирование")
    add_paragraph(doc, "Тестирование проводилось вручную по сценариям:")
    add_numbered(
        doc,
        "Регистрация → создание поездки → добавление точки → проверка "
        "на карте → расход → лайк → комментарий другим пользователем.",
    )
    add_numbered(
        doc,
        "API-проверки через curl: GET /api/trips, POST без авторизации "
        "(403), GET для чужой приватной поездки (403), геокодер для "
        "адреса «Москва».",
    )
    add_numbered(
        doc,
        "Граничные случаи: пустой бюджет, точка без координат, поездка "
        "на 1 день, поездка на 300+ дней.",
    )
    add_numbered(
        doc,
        "Безопасность: попытка отправить форму без CSRF-токена → 400.",
    )

    add_h1(doc, "6. Заключение")

    add_h2(doc, "Что получилось")
    add_bullet(doc, "Полноценное Flask-приложение из 4 410 строк кода.")
    add_bullet(doc, "Закрыты все запрошенные технологии курса.")
    add_bullet(
        doc,
        "Реализован уникальный сценарий — интерактивная карта с "
        "построением маршрута и геокодером.",
    )
    add_bullet(doc, "Проект развёрнут в облаке (Render.com).")

    add_h2(doc, "Что можно доработать")
    add_bullet(
        doc,
        "Импорт и экспорт маршрутов в формат GPX/KML для офлайн-карт.",
    )
    add_bullet(
        doc,
        "Совместная редактура поездки (несколько владельцев).",
    )
    add_bullet(
        doc,
        "Уведомления (письма на e-mail) о новых комментариях.",
    )
    add_bullet(
        doc,
        "Кэширование тайлов и геокодера для снижения нагрузки на "
        "сторонние сервисы.",
    )
    add_bullet(
        doc,
        "Покрытие тестами на pytest (сейчас — ручное тестирование).",
    )
    add_bullet(doc, "Мобильное приложение, использующее REST API.")

    add_h2(doc, "Чему научился")
    add_bullet(
        doc,
        "Проектировать многослойные Flask-приложения с фабрикой.",
    )
    add_bullet(
        doc,
        "Использовать SQLAlchemy 2.0 с типизированными `Mapped[...]`.",
    )
    add_bullet(doc, "Делать AJAX-взаимодействие на чистом JS без jQuery.")
    add_bullet(
        doc,
        "Готовить Flask-проект к production-деплою (gunicorn, ENV-vars, "
        "разные БД для dev/prod).",
    )
    add_bullet(
        doc,
        "Документировать проект и оформлять README, ТЗ и пояснительную "
        "записку.",
    )

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    print(f"Пояснительная записка сохранена: {target}")


# --- Сборка ------------------------------------------------------------------

def build_all() -> None:
    project_root = Path(__file__).resolve().parent.parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))

    docs_dir = project_root / "docs"
    author = Author()

    build_tz(docs_dir / "ТЗ.docx", author)
    build_explanatory_note(
        docs_dir / "Пояснительная_записка.docx", author,
    )


if __name__ == "__main__":
    build_all()
